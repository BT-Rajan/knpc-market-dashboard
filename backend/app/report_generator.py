"""
Quarterly market report generation for MOG division.
Compiles price movements, product trends, and market developments into Word documents.
"""
from datetime import datetime, date
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from sqlalchemy.orm import Session

from app.models import Item, PriceHistory, NewsItem
from app.config import QUARTER_MONTHS, MOG_DIVISION_NAME, REPORTS_DIR

GOLD = RGBColor(0xB8, 0x8A, 0x1E)
DARK = RGBColor(0x1A, 0x1F, 0x2B)
GREY = RGBColor(0x5A, 0x66, 0x78)


def _set_cell_text(cell, text, bold=False, color=None, size=10):
    """Format a table cell with text."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _add_table(doc, headers, rows):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, color=GOLD)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            _set_cell_text(cells[i], "" if val is None else val)
    return table


def _heading(doc, text, level=1):
    """Add a formatted heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK
    return h


def get_quarter_date_range(year: int, quarter: str) -> tuple[date, date]:
    """Get start and end dates for a quarter."""
    if quarter not in QUARTER_MONTHS:
        raise ValueError(f"Invalid quarter: {quarter}")
    
    start_month, end_month = QUARTER_MONTHS[quarter]
    start = date(year, start_month, 1)
    
    if end_month == 12:
        end = date(year, 12, 31)
    else:
        end = date(year, end_month + 1, 1)
        end = date(end.year, end.month, end.day - 1)  # Last day of end_month
    
    return start, end


def get_benchmark_stats(db: Session, year: int, quarter: str) -> list[dict]:
    """Calculate quarterly stats for crude benchmarks."""
    start, end = get_quarter_date_range(year, quarter)
    
    benchmarks = ["BRENT", "WTI", "OMAN", "DUBAI", "KEC"]
    stats = []
    
    for code in benchmarks:
        item = db.query(Item).filter(Item.code == code).first()
        if not item:
            continue
            
        prices = (
            db.query(PriceHistory)
            .filter(
                PriceHistory.item_id == item.id,
                PriceHistory.price_date >= start,
                PriceHistory.price_date <= end
            )
            .order_by(PriceHistory.price_date)
            .all()
        )
        
        if not prices:
            continue
        
        prices_list = [p.price for p in prices]
        opening = prices_list[0]
        closing = prices_list[-1]
        high = max(prices_list)
        low = min(prices_list)
        avg = sum(prices_list) / len(prices_list)
        change = closing - opening
        change_pct = (change / opening * 100) if opening != 0 else 0
        
        stats.append({
            "code": code,
            "name": item.name,
            "open": round(opening, 2),
            "close": round(closing, 2),
            "high": round(high, 2),
            "low": round(low, 2),
            "avg": round(avg, 2),
            "change": round(change, 2),
            "change_pct": round(change_pct, 2),
            "readings": len(prices)
        })
    
    return stats


def get_product_stats(db: Session, year: int, quarter: str) -> list[dict]:
    """Calculate quarterly stats for refined products."""
    start, end = get_quarter_date_range(year, quarter)
    
    products = ["NAPHTHA", "GASOLINE92", "GASOLINE95", "JETKERO", "GASOIL10", "FUELOIL180", "FUELOIL380", "LPG"]
    stats = []
    
    for code in products:
        item = db.query(Item).filter(Item.code == code).first()
        if not item:
            continue
            
        prices = (
            db.query(PriceHistory)
            .filter(
                PriceHistory.item_id == item.id,
                PriceHistory.price_date >= start,
                PriceHistory.price_date <= end
            )
            .order_by(PriceHistory.price_date)
            .all()
        )
        
        if not prices:
            continue
        
        prices_list = [p.price for p in prices]
        opening = prices_list[0]
        closing = prices_list[-1]
        high = max(prices_list)
        low = min(prices_list)
        avg = sum(prices_list) / len(prices_list)
        change = closing - opening
        change_pct = (change / opening * 100) if opening != 0 else 0
        
        stats.append({
            "code": code,
            "name": item.name,
            "open": round(opening, 2),
            "close": round(closing, 2),
            "high": round(high, 2),
            "low": round(low, 2),
            "avg": round(avg, 2),
            "change": round(change, 2),
            "change_pct": round(change_pct, 2),
            "readings": len(prices)
        })
    
    return stats


def get_recent_news(db: Session, limit: int = 20) -> list[dict]:
    """Get recent news items."""
    items = (
        db.query(NewsItem)
        .order_by(NewsItem.collected_at.desc())
        .limit(limit)
        .all()
    )
    return [
        {
            "headline": n.headline,
            "url": n.url,
            "source": n.source,
            "collected_at": n.collected_at.strftime("%d %b %Y") if n.collected_at else "Unknown"
        }
        for n in items
    ]


def generate_quarterly_report(
    db: Session,
    year: int,
    quarter: str,
    outlook_notes: str = "",
    generated_by: str = "MOG Analyst"
) -> bytes:
    """
    Generate a quarterly report as a Word document.
    Returns the document as bytes (for download).
    """
    b_stats = get_benchmark_stats(db, year, quarter)
    p_stats = get_product_stats(db, year, quarter)
    news_items = get_recent_news(db, limit=20)
    
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{quarter} {year} Market Intelligence Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = DARK
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(MOG_DIVISION_NAME)
    run.font.size = Pt(13)
    run.font.color.rgb = GOLD
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated {datetime.now().strftime('%d %B %Y')} • Prepared by {generated_by}")
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    
    doc.add_paragraph()
    
    # Executive summary
    _heading(doc, "Executive Summary", level=1)
    if b_stats:
        sorted_b = sorted(b_stats, key=lambda x: x["change_pct"], reverse=True)
        best = sorted_b[0]
        worst = sorted_b[-1]
        summary_text = (
            f"During {quarter} {year}, {best['name']} was the strongest-performing tracked "
            f"benchmark ({best['change_pct']:+.2f}% quarter-on-quarter), "
            f"while {worst['name']} moved {worst['change_pct']:+.2f}%. "
            f"{len(news_items)} market developments were logged in the monitoring feed."
        )
    else:
        summary_text = (
            f"No benchmark price readings were recorded for {quarter} {year} at the time of "
            f"generation. Populate the pipeline or enter manual readings, then regenerate this report."
        )
    doc.add_paragraph(summary_text)
    
    # Section I: Crude benchmark review
    _heading(doc, "I. Crude Benchmark Price Review", level=1)
    doc.add_paragraph(
        "Quarter-over-quarter movement in Brent, WTI, Dubai, Oman, and Kuwait Export Crude, "
        "based on daily readings logged in the market intelligence system."
    )
    if b_stats:
        _add_table(
            doc,
            ["Benchmark", "Open", "Close", "High", "Low", "Average", "Change", "Change %", "Readings"],
            [[
                s["name"],
                s["open"],
                s["close"],
                s["high"],
                s["low"],
                s["avg"],
                f"{s['change']:+.2f}",
                f"{s['change_pct']:+.2f}%",
                s["readings"]
            ] for s in b_stats],
        )
    else:
        doc.add_paragraph("No data available for this period.")
    
    doc.add_paragraph()
    
    # Section II: Refined products
    _heading(doc, "II. Refined Product Proxy Review", level=1)
    doc.add_paragraph(
        "Singapore / regional refined product proxy movement (Naphtha, Gasoline, Jet/Kerosene, "
        "Gasoil, Fuel Oil, LPG)."
    )
    if p_stats:
        _add_table(
            doc,
            ["Product", "Open", "Close", "High", "Low", "Average", "Change", "Change %", "Readings"],
            [[
                s["name"],
                s["open"],
                s["close"],
                s["high"],
                s["low"],
                s["avg"],
                f"{s['change']:+.2f}",
                f"{s['change_pct']:+.2f}%",
                s["readings"]
            ] for s in p_stats],
        )
    else:
        doc.add_paragraph("No data available for this period.")
    
    doc.add_paragraph()
    
    # Section III: Market Intelligence
    _heading(doc, "III. Market Intelligence & News", level=1)
    doc.add_paragraph("Recent market developments and news items logged in the monitoring feed.")
    if news_items:
        for item in news_items[:10]:  # Top 10 news items
            doc.add_paragraph(
                f"{item['headline']} ({item['source']}, {item['collected_at']})",
                style="List Bullet"
            )
    else:
        doc.add_paragraph("No news items logged for this period.")
    
    doc.add_paragraph()
    
    # Outlook
    if outlook_notes.strip():
        _heading(doc, "IV. Analyst Outlook & Commentary", level=1)
        doc.add_paragraph(outlook_notes)
    
    # Convert to bytes
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def save_quarterly_report(
    db: Session,
    year: int,
    quarter: str,
    outlook_notes: str = "",
    generated_by: str = "MOG Analyst"
) -> Path:
    """
    Generate and save a quarterly report to disk.
    Returns the file path.
    """
    report_bytes = generate_quarterly_report(db, year, quarter, outlook_notes, generated_by)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{quarter}_{year}_market_report_{timestamp}.docx"
    filepath = REPORTS_DIR / filename
    
    with open(filepath, "wb") as f:
        f.write(report_bytes)
    
    return filepath


def list_generated_reports() -> list[Path]:
    """List all generated quarterly reports."""
    return sorted(REPORTS_DIR.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
